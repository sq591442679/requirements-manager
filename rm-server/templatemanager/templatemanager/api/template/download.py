import base64
import io

import docx
from flask import request

from templatemanager.app import app
from templatemanager.utils.handle_api import handle_response

META_SUCCESS = {'status': 200, 'msg': '模板生成成功！'}
META_ERROR_NOT_EXIST = {'status': 404, 'msg': '生成失败，该模板不存在！'}


# @app.route('/template/download', methods=['GET'])
# @handle_download
# def template_download():
#     body = request.args.get('file')
#
#     template_name = body
#     print('name:' + template_name)
#     template_mongodb_dao = TemplateMongoDBDao(template_collection)
#     # check if template not exists
#     template = template_mongodb_dao.get_template(template_name)
#     if not template:
#         return {'meta': META_ERROR_NOT_EXIST}
#
#     template_word_docx = Document()
#     template_word_docx.add_heading(template.template_name, 0)
#     for para in template.outline:
#         template_word_docx.add_paragraph(para)
#
#     # save the file to a file stream
#     docx_file = io.FileIO()
#     template_word_docx.save(docx_file)
#
#     ret_info = {
#         'meta': META_SUCCESS,
#         'file': docx_file
#     }
#     docx_file.close()
#
#     return {
#         'meta': META_SUCCESS,
#     }
@app.route('/template/download', methods=['GET'])
@handle_response
def template_download():
    body = request.args.get('file')
    print('body:' + body)
    document = docx.Document('E:\\BUAA\\BUAA课程\\大三下\\生产实习\\需求管理工具\\RequirementsManager-master'
                             '\\rm-server\\templatemanager\\templatemanager\\uploads\\' + body + '.docx')
    for p in document.paragraphs:
        print(p.text)
    docx_file = io.BytesIO()
    document.save(docx_file)

    return {
        'data': {
            'file_base64': base64.b64encode(docx_file.getvalue()).decode('utf-8'),
        },
        'meta': META_SUCCESS
    }
